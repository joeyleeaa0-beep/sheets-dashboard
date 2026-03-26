import streamlit as st
import pandas as pd
import requests
import plotly.express as px
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime
import os

st.set_page_config(page_title="新媒体数据看板", page_icon="📊", layout="wide")

st.markdown("""
<style>
    .stApp { background-color: #f8f9fc; }
    #MainMenu, footer, header { visibility: hidden; }
    [data-testid="metric-container"] {
        background: white;
        border: 1px solid #eef0f4;
        border-radius: 12px;
        padding: 20px 24px;
        box-shadow: 0 1px 4px rgba(0,0,0,0.05);
    }
    [data-testid="metric-container"] label { color: #6b7280 !important; font-size: 13px !important; font-weight: 500 !important; }
    [data-testid="metric-container"] [data-testid="stMetricValue"] { color: #111827 !important; font-size: 26px !important; font-weight: 700 !important; }
    .stTabs [data-baseweb="tab-list"] { gap: 8px; background: transparent; border-bottom: 1px solid #eef0f4; }
    .stTabs [data-baseweb="tab"] { background: transparent; color: #6b7280; font-weight: 500; border-radius: 6px 6px 0 0; padding: 8px 18px; }
    .stTabs [aria-selected="true"] { background: white !important; color: #111827 !important; border-bottom: 2px solid #4f46e5 !important; }
    [data-testid="stSidebar"] { background: white; border-right: 1px solid #eef0f4; }
    hr { border-color: #eef0f4; }
</style>
""", unsafe_allow_html=True)

APP_ID = st.secrets["FEISHU_APP_ID"]
APP_SECRET = st.secrets["FEISHU_APP_SECRET"]
WIKI_TOKEN = "G0VbwCGGMiEVNlktEYwcoA4JnHd"
SHEETS = {"四地汇总": "0HBHQk", "细分项汇总": "1hOmjY"}
MONTHS = ["1月","2月","3月","4月","5月","6月","7月","8月","9月","10月","11月","12月"]
CITIES = ["深圳", "上海", "成都", "天津"]
COLORS = ["#4f46e5","#06b6d4","#10b981","#f59e0b","#ef4444","#8b5cf6"]
PLOTLY_LAYOUT = dict(
    paper_bgcolor="white", plot_bgcolor="white",
    font=dict(family="sans-serif", size=12, color="#374151"),
    margin=dict(l=16, r=16, t=48, b=16),
    legend=dict(
        orientation="h",
        yanchor="top", y=-0.2,
        xanchor="center", x=0.5,
        font=dict(size=11),
    ),
    xaxis=dict(showgrid=False, linecolor="#eef0f4"),
    yaxis=dict(gridcolor="#f3f4f6", linecolor="#eef0f4"),
)

# ── 中文字体设置 ──
@st.cache_resource
def setup_chinese_font():
    font_path = "/tmp/NotoSansSC.ttf"
    if not os.path.exists(font_path):
        urls = [
            "https://cdn.jsdelivr.net/gh/googlefonts/noto-cjk@main/Sans/SubsetOTF/SC/NotoSansSC-Regular.otf",
            "https://github.com/googlefonts/noto-cjk/raw/main/Sans/SubsetOTF/SC/NotoSansSC-Regular.otf",
        ]
        for url in urls:
            try:
                r = requests.get(url, timeout=30)
                if r.status_code == 200:
                    with open(font_path, "wb") as f:
                        f.write(r.content)
                    break
            except Exception:
                continue
    if os.path.exists(font_path):
        try:
            fm.fontManager.addfont(font_path)
            prop = fm.FontProperties(fname=font_path)
            return prop.get_name(), font_path
        except Exception:
            pass
    return None, None

CHINESE_FONT_NAME, CHINESE_FONT_PATH = setup_chinese_font()
CHINESE_FONT = CHINESE_FONT_NAME or "DejaVu Sans"
plt.rcParams["font.sans-serif"] = [CHINESE_FONT, "DejaVu Sans"]
plt.rcParams["axes.unicode_minus"] = False

def get_font_prop():
    if CHINESE_FONT_PATH and os.path.exists(CHINESE_FONT_PATH):
        return fm.FontProperties(fname=CHINESE_FONT_PATH)
    return fm.FontProperties(family=CHINESE_FONT)

# ── API ──
@st.cache_data(ttl=60)
def get_token():
    res = requests.post(
        "https://open.feishu.cn/open-apis/auth/v3/tenant_access_token/internal",
        json={"app_id": APP_ID, "app_secret": APP_SECRET}
    )
    return res.json().get("tenant_access_token")

@st.cache_data(ttl=60)
def get_spreadsheet_token():
    token = get_token()
    headers = {"Authorization": f"Bearer {token}"}
    res = requests.get(
        f"https://open.feishu.cn/open-apis/wiki/v2/spaces/get_node?token={WIKI_TOKEN}",
        headers=headers
    ).json()
    return res.get("data", {}).get("node", {}).get("obj_token")

@st.cache_data(ttl=60)
def read_sheet(spreadsheet_token, sheet_id, data_range="A1:Z500"):
    token = get_token()
    headers = {"Authorization": f"Bearer {token}"}
    url = f"https://open.feishu.cn/open-apis/sheets/v2/spreadsheets/{spreadsheet_token}/values/{sheet_id}!{data_range}?renderType=FORMATTED_VALUE"
    res = requests.get(url, headers=headers).json()
    values = res.get("data", {}).get("valueRange", {}).get("values", [])
    if not values or len(values) < 2:
        return pd.DataFrame()
    headers_row = [str(h) if h else f"列{i}" for i, h in enumerate(values[0])]
    return pd.DataFrame(values[1:], columns=headers_row)

# ── 工具函数 ──
def to_num(series):
    if isinstance(series, pd.Series):
        s = series.astype(str)
        s = s.str.replace(r'IFERROR\(.*?\)', '0', regex=True)
        s = s.str.replace("/", "0").str.replace(",", "")
        return pd.to_numeric(s, errors="coerce").fillna(0)
    return 0

def fill_merged(df, col):
    if col in df.columns:
        df[col] = df[col].replace([None, "None", "", "nan", "/"], pd.NA)
        df[col] = df[col].ffill()
    return df

def apply_filter(df, city, month):
    if df is None or df.empty:
        return pd.DataFrame()
    d = df.copy()
    if city != "全部城市" and "地区" in d.columns:
        d = d[d["地区"] == city]
    if month != "全部月份" and "月份" in d.columns:
        d = d[d["月份"] == month]
    return d

def classify_channel(x):
    x = str(x)
    if "抖音" in x: return "抖音"
    if "信息流" in x: return "信息流"
    if "小红书" in x: return "小红书"
    if "视频号" in x: return "视频号"
    if "B站" in x: return "B站"
    if "快手" in x: return "快手"
    return x

def make_chart(fig):
    fig.update_layout(**PLOTLY_LAYOUT)
    return fig

# ── 数据清洗 ──
def clean_main_df(df):
    if df.empty: return df
    df = fill_merged(df, "月份")
    if "月份" in df.columns:
        df["月份"] = df["月份"].astype(str).str.strip().str.replace(" ", "")
    if "地区" in df.columns:
        df["地区"] = df["地区"].replace([None,"None","","nan"], pd.NA)
        df = df[df["地区"].isin(CITIES)].copy()
    for col in ["投放金额","总成交量","销售量","收购量"]:
        if col in df.columns:
            df[col] = to_num(df[col])
    for col in df.columns:
        if "客资" in col and "成本" not in col:
            df["客资量"] = to_num(df[col])
            break
    if "投放金额" in df.columns and "客资量" in df.columns:
        df["综合客资成本"] = (df["投放金额"] / df["客资量"].replace(0, pd.NA)).round(2)
    if "投放金额" in df.columns and "总成交量" in df.columns:
        df["综合成交成本"] = (df["投放金额"] / df["总成交量"].replace(0, pd.NA)).round(2)
    return df

def clean_detail_df(df):
    if df.empty: return pd.DataFrame()
    header_idx = 0
    for i in range(min(5, len(df))):
        row_vals = [str(v) for v in df.iloc[i].values]
        if any("渠道" in v or "月份" in v for v in row_vals):
            header_idx = i
            break
    new_cols = [str(v).strip() if v and str(v) != "None" else f"列{i}"
                for i, v in enumerate(df.iloc[header_idx].values)]
    df = df.iloc[header_idx+1:].reset_index(drop=True)
    df.columns = new_cols
    df = fill_merged(df, "月份")
    df = fill_merged(df, "地区")
    if "月份" in df.columns:
        df["月份"] = df["月份"].astype(str).str.strip().str.replace(" ", "")
    channel_col = next((c for c in df.columns if "渠道" in c or "平台" in c), None)
    if channel_col:
        df = df[~df[channel_col].astype(str).str.contains("合计", na=False)]
        df = df[df[channel_col].astype(str).str.strip().isin(["","None","nan"]) == False]
        df = df[df[channel_col].notna()].copy()
        df["渠道分类"] = df[channel_col].apply(classify_channel)
    if "地区" in df.columns:
        df = df[df["地区"].isin(CITIES)].copy()
    for col in df.columns:
        if col in ["投放金额","销售量","收购量","素材更新量","直播场次"]:
            df[col] = to_num(df[col])
        elif "客资" in col and "成本" not in col and "直播" not in col and "视频" not in col:
            df["客资量"] = to_num(df[col])
    if "销售量" in df.columns and "收购量" in df.columns:
        df["总成交量"] = df["销售量"] + df["收购量"]
    return df

# ── matplotlib 图表 ──
def bar_image(df, x_col, y_col, title, width=8, height=3.5):
    fp = get_font_prop()
    fig, ax = plt.subplots(figsize=(width, height))
    vals = df[y_col].tolist()
    labels = df[x_col].astype(str).tolist()
    bar_colors = [COLORS[i % len(COLORS)] for i in range(len(labels))]
    bars = ax.bar(labels, vals, color=bar_colors, edgecolor='white', linewidth=0.5, width=0.5)
    ax.set_title(title, fontsize=12, pad=12, color='#111827', fontweight='bold', fontproperties=fp)
    for spine in ['top','right']:
        ax.spines[spine].set_visible(False)
    for spine in ['left','bottom']:
        ax.spines[spine].set_color('#e5e7eb')
    ax.tick_params(colors='#6b7280', labelsize=9)
    ax.yaxis.grid(True, color='#f3f4f6', linewidth=0.8)
    ax.set_axisbelow(True)
    fig.patch.set_facecolor('white')
    ax.set_facecolor('white')
    for bar, v in zip(bars, vals):
        if v > 0:
            ax.text(bar.get_x() + bar.get_width()/2., v * 1.01,
                    f'{v:,.0f}', ha='center', va='bottom', fontsize=8, color='#374151')
    for label in ax.get_xticklabels():
        label.set_font_properties(fp)
    plt.tight_layout()
    buf = BytesIO()
    plt.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    buf.seek(0)
    return buf

def line_image(df, x_col, y_col, color_col, title, width=8, height=3.5):
    fp = get_font_prop()
    fig, ax = plt.subplots(figsize=(width, height))
    groups = df[color_col].unique()
    for i, grp in enumerate(groups):
        d = df[df[color_col] == grp].copy()
        ax.plot(d[x_col].astype(str), d[y_col],
                marker='o', color=COLORS[i % len(COLORS)],
                linewidth=2, markersize=5, label=str(grp))
    ax.set_title(title, fontsize=12, pad=12, color='#111827', fontweight='bold', fontproperties=fp)
    for spine in ['top','right']:
        ax.spines[spine].set_visible(False)
    for spine in ['left','bottom']:
        ax.spines[spine].set_color('#e5e7eb')
    ax.tick_params(colors='#6b7280', labelsize=9)
    ax.yaxis.grid(True, color='#f3f4f6', linewidth=0.8)
    ax.set_axisbelow(True)
    ax.legend(fontsize=8, framealpha=0, loc='upper left', prop=fp)
    fig.patch.set_facecolor('white')
    ax.set_facecolor('white')
    for label in ax.get_xticklabels():
        label.set_font_properties(fp)
    plt.tight_layout()
    buf = BytesIO()
    plt.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    buf.seek(0)
    return buf

# ── Word 文档生成 ──
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_section_title(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x11, 0x18, 0x27)
    return p

def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'EEF0F4')
    pBdr.append(bottom)
    pPr.append(pBdr)

def generate_word(sel_city, sel_month, metrics, cg, rg, cm_df, ch_month_df):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("新媒体数据看板报告")
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x11, 0x18, 0x27)

    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(16)
    run2 = sub.add_run(
        f"城市：{sel_city}  |  月份：{sel_month}  |  生成时间：{datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}"
    )
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(0x6b, 0x72, 0x80)

    add_divider(doc)

    add_section_title(doc, "📊 核心指标总览")
    m = metrics
    metric_rows = [
        [("总投放金额", m['total_spend']), ("总客资量", m['total_keizi']),
         ("总成交量", m['total_chengjiao']), ("成交率", m['chengjiao_rate'])],
        [("销售总量", m['total_xiaoshou']), ("收购总量", m['total_shougou']),
         ("客资成本", m['keizi_cost']), ("成交成本", m['chengjiao_cost'])],
    ]
    for row_data in metric_rows:
        table = doc.add_table(rows=2, cols=4)
        table.style = 'Table Grid'
        table.autofit = False
        col_w = Cm(4.25)
        for i, (label, value) in enumerate(row_data):
            label_cell = table.rows[0].cells[i]
            label_cell.width = col_w
            label_cell.paragraphs[0].clear()
            label_p = label_cell.paragraphs[0]
            label_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            lr = label_p.add_run(label)
            lr.font.size = Pt(9)
            lr.font.color.rgb = RGBColor(0x6b, 0x72, 0x80)
            set_cell_bg(label_cell, 'F3F4F6')
            val_cell = table.rows[1].cells[i]
            val_cell.width = col_w
            val_cell.paragraphs[0].clear()
            val_p = val_cell.paragraphs[0]
            val_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            vr = val_p.add_run(str(value))
            vr.font.size = Pt(16)
            vr.font.bold = True
            vr.font.color.rgb = RGBColor(0x11, 0x18, 0x27)
            set_cell_bg(val_cell, 'FFFFFF')
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    add_divider(doc)

    if cg is not None and not cg.empty:
        add_section_title(doc, "🏙️ 分城市经营对比")
        city_cols = [c for c in ["地区","投放金额","客资量","总成交量","客资成本","成交成本","成交率%"] if c in cg.columns]
        t = doc.add_table(rows=1+len(cg), cols=len(city_cols))
        t.style = 'Table Grid'
        for i, col in enumerate(city_cols):
            cell = t.rows[0].cells[i]
            cell.paragraphs[0].clear()
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(col)
            r.font.size = Pt(9)
            r.font.bold = True
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            set_cell_bg(cell, '4F46E5')
        for ri, (_, row) in enumerate(cg[city_cols].iterrows()):
            bg = 'FFFFFF' if ri % 2 == 0 else 'F9FAFB'
            for ci, col in enumerate(city_cols):
                cell = t.rows[ri+1].cells[ci]
                cell.paragraphs[0].clear()
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                v = row[col]
                text = str(round(v, 2)) if isinstance(v, float) else str(v)
                r = p.add_run(text)
                r.font.size = Pt(9)
                set_cell_bg(cell, bg)
        doc.add_paragraph().paragraph_format.space_after = Pt(8)
        img = bar_image(cg, "地区", "客资量", "各城市客资量对比")
        doc.add_picture(img, width=Cm(17))
        doc.add_paragraph().paragraph_format.space_after = Pt(4)
        img2 = bar_image(cg, "地区", "客资成本", "各城市客资成本对比")
        doc.add_picture(img2, width=Cm(17))

    add_divider(doc)

    if rg is not None and not rg.empty:
        add_section_title(doc, "📡 分渠道数据对比")
        ch_cols = [c for c in ["渠道分类","投放金额","客资量","总成交量","客资成本"] if c in rg.columns]
        t2 = doc.add_table(rows=1+len(rg), cols=len(ch_cols))
        t2.style = 'Table Grid'
        for i, col in enumerate(ch_cols):
            cell = t2.rows[0].cells[i]
            cell.paragraphs[0].clear()
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(col)
            r.font.size = Pt(9)
            r.font.bold = True
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            set_cell_bg(cell, '4F46E5')
        for ri, (_, row) in enumerate(rg[ch_cols].iterrows()):
            bg = 'FFFFFF' if ri % 2 == 0 else 'F9FAFB'
            for ci, col in enumerate(ch_cols):
                cell = t2.rows[ri+1].cells[ci]
                cell.paragraphs[0].clear()
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                v = row[col]
                text = str(round(v, 2)) if isinstance(v, float) else str(v)
                r = p.add_run(text)
                r.font.size = Pt(9)
                set_cell_bg(cell, bg)
        doc.add_paragraph().paragraph_format.space_after = Pt(8)
        img3 = bar_image(rg, "渠道分类", "客资量", "各渠道客资量对比")
        doc.add_picture(img3, width=Cm(17))
        doc.add_paragraph().paragraph_format.space_after = Pt(4)
        img4 = bar_image(rg, "渠道分类", "投放金额", "各渠道投放金额对比")
        doc.add_picture(img4, width=Cm(17))

    add_divider(doc)

    if cm_df is not None and not cm_df.empty:
        add_section_title(doc, "📈 趋势分析")
        img5 = line_image(cm_df, "月份", "客资成本", "地区", "各城市客资成本月度趋势")
        doc.add_picture(img5, width=Cm(17))
        doc.add_paragraph().paragraph_format.space_after = Pt(4)
        img6 = line_image(cm_df, "月份", "成交成本", "地区", "各城市成交成本月度趋势")
        doc.add_picture(img6, width=Cm(17))

    if ch_month_df is not None and not ch_month_df.empty:
        doc.add_paragraph().paragraph_format.space_after = Pt(4)
        img7 = line_image(ch_month_df, "月份", "客资量", "渠道分类", "各渠道客资量月度趋势")
        doc.add_picture(img7, width=Cm(17))

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()

# ── 加载数据 ──
with st.spinner("正在加载数据..."):
    try:
        spreadsheet_token = get_spreadsheet_token()
        if not spreadsheet_token:
            st.error("无法获取表格Token")
            st.stop()
        df_main = clean_main_df(read_sheet(spreadsheet_token, SHEETS["四地汇总"], "A2:K200").copy())
        df_detail = clean_detail_df(read_sheet(spreadsheet_token, SHEETS["细分项汇总"], "A1:P500").copy())
        if df_detail is None: df_detail = pd.DataFrame()
    except Exception as e:
        st.error(f"加载失败：{e}")
        st.stop()

# ── 侧边栏 ──
with st.sidebar:
    st.markdown("## 筛选条件")
    sel_city = st.selectbox("城市", ["全部城市"] + CITIES)
    sel_month = st.selectbox("月份", ["全部月份"] + MONTHS)
    st.divider()
    df_filtered = apply_filter(df_main, sel_city, sel_month)
    df_detail_filtered = apply_filter(df_detail, sel_city, sel_month)
    st.caption(f"四地汇总：{len(df_filtered)} 条")
    st.caption(f"渠道明细：{len(df_detail_filtered)} 条")
    st.divider()
    st.markdown("## 导出报告")
    if st.button("生成 Word 报告", use_container_width=True):
        with st.spinner("正在生成报告..."):
            total_spend = to_num(df_filtered["投放金额"]).sum() if "投放金额" in df_filtered.columns else 0
            total_keizi = to_num(df_filtered["客资量"]).sum() if "客资量" in df_filtered.columns else 0
            total_chengjiao = to_num(df_filtered["总成交量"]).sum() if "总成交量" in df_filtered.columns else 0
            total_xiaoshou = to_num(df_filtered["销售量"]).sum() if "销售量" in df_filtered.columns else 0
            total_shougou = to_num(df_filtered["收购量"]).sum() if "收购量" in df_filtered.columns else 0
            keizi_cost = total_spend / total_keizi if total_keizi > 0 else 0
            chengjiao_cost = total_spend / total_chengjiao if total_chengjiao > 0 else 0
            chengjiao_rate = total_chengjiao / total_keizi * 100 if total_keizi > 0 else 0
            metrics = {
                'total_spend': f"¥{total_spend:,.0f}",
                'total_keizi': f"{int(total_keizi):,}",
                'total_chengjiao': f"{int(total_chengjiao):,}",
                'chengjiao_rate': f"{chengjiao_rate:.2f}%",
                'total_xiaoshou': f"{int(total_xiaoshou):,}",
                'total_shougou': f"{int(total_shougou):,}",
                'keizi_cost': f"¥{keizi_cost:.2f}",
                'chengjiao_cost': f"¥{chengjiao_cost:.2f}",
            }
            df_city = apply_filter(df_main, "全部城市", sel_month)
            cg = None
            if not df_city.empty and "地区" in df_city.columns:
                cg = df_city.groupby("地区").agg(投放金额=("投放金额","sum"), 客资量=("客资量","sum"), 总成交量=("总成交量","sum")).reset_index()
                cg["客资成本"] = (cg["投放金额"]/cg["客资量"].replace(0,pd.NA)).round(2)
                cg["成交成本"] = (cg["投放金额"]/cg["总成交量"].replace(0,pd.NA)).round(2)
                cg["成交率%"] = (cg["总成交量"]/cg["客资量"].replace(0,pd.NA)*100).round(2)
            rg = None
            if not df_detail_filtered.empty and "渠道分类" in df_detail_filtered.columns:
                rg = df_detail_filtered.groupby("渠道分类").agg(投放金额=("投放金额","sum"), 客资量=("客资量","sum"), 总成交量=("总成交量","sum")).reset_index()
                rg["客资成本"] = (rg["投放金额"]/rg["客资量"].replace(0,pd.NA)).round(2)
                rg = rg.sort_values("客资量", ascending=False)
            cm_df = None
            if not df_main.empty and "月份" in df_main.columns:
                df_t = df_main.copy()
                if sel_city != "全部城市" and "地区" in df_t.columns:
                    df_t = df_t[df_t["地区"] == sel_city]
                cm_df = df_t.groupby(["月份","地区"]).agg(投放金额=("投放金额","sum"), 客资量=("客资量","sum"), 总成交量=("总成交量","sum")).reset_index()
                cm_df["客资成本"] = (cm_df["投放金额"]/cm_df["客资量"].replace(0,pd.NA)).round(2)
                cm_df["成交成本"] = (cm_df["投放金额"]/cm_df["总成交量"].replace(0,pd.NA)).round(2)
                cm_df["月份"] = pd.Categorical(cm_df["月份"], categories=MONTHS, ordered=True)
                cm_df = cm_df.sort_values("月份")
            ch_month_df = None
            if not df_detail.empty and "月份" in df_detail.columns and "渠道分类" in df_detail.columns:
                df_ch = apply_filter(df_detail, sel_city, "全部月份")
                ch_month_df = df_ch.groupby(["月份","渠道分类"]).agg(客资量=("客资量","sum")).reset_index()
                ch_month_df["月份"] = pd.Categorical(ch_month_df["月份"], categories=MONTHS, ordered=True)
                ch_month_df = ch_month_df.sort_values("月份")
            try:
                word_bytes = generate_word(sel_city, sel_month, metrics, cg, rg, cm_df, ch_month_df)
                st.download_button(
                    label="⬇️ 下载 Word 报告",
                    data=word_bytes,
                    file_name=f"新媒体看板报告_{sel_city}_{sel_month}_{datetime.datetime.now().strftime('%Y%m%d')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
            except Exception as e:
                st.error(f"报告生成失败：{e}")

# ── 主页面 ──
st.markdown(f"""
<div style="padding:8px 0 20px 0;border-bottom:1px solid #eef0f4;margin-bottom:24px;">
    <h2 style="margin:0;color:#111827;font-weight:700;">📊 新媒体数据看板</h2>
    <p style="margin:4px 0 0 0;color:#6b7280;font-size:14px;">{sel_city} · {sel_month} · 数据每60秒自动更新</p>
</div>
""", unsafe_allow_html=True)

total_spend = to_num(df_filtered["投放金额"]).sum() if "投放金额" in df_filtered.columns else 0
total_keizi = to_num(df_filtered["客资量"]).sum() if "客资量" in df_filtered.columns else 0
total_chengjiao = to_num(df_filtered["总成交量"]).sum() if "总成交量" in df_filtered.columns else 0
total_xiaoshou = to_num(df_filtered["销售量"]).sum() if "销售量" in df_filtered.columns else 0
total_shougou = to_num(df_filtered["收购量"]).sum() if "收购量" in df_filtered.columns else 0
keizi_cost = total_spend / total_keizi if total_keizi > 0 else 0
chengjiao_cost = total_spend / total_chengjiao if total_chengjiao > 0 else 0
chengjiao_rate = total_chengjiao / total_keizi * 100 if total_keizi > 0 else 0

c1,c2,c3,c4 = st.columns(4)
c1.metric("总投放金额", f"¥{total_spend:,.0f}")
c2.metric("总客资量", f"{int(total_keizi):,}")
c3.metric("总成交量", f"{int(total_chengjiao):,}")
c4.metric("成交率", f"{chengjiao_rate:.2f}%")
c5,c6,c7,c8 = st.columns(4)
c5.metric("销售总量", f"{int(total_xiaoshou):,}")
c6.metric("收购总量", f"{int(total_shougou):,}")
c7.metric("客资成本", f"¥{keizi_cost:.2f}")
c8.metric("成交成本", f"¥{chengjiao_cost:.2f}")

st.divider()

tab1, tab2, tab3, tab4 = st.tabs(["🏙️ 分城市", "📡 分渠道", "📈 趋势分析", "📋 数据明细"])

with tab1:
    st.subheader("分城市经营对比")
    df_city = apply_filter(df_main, "全部城市", sel_month)
    if not df_city.empty and "地区" in df_city.columns:
        cg = df_city.groupby("地区").agg(投放金额=("投放金额","sum"), 客资量=("客资量","sum"), 总成交量=("总成交量","sum")).reset_index()
        cg["客资成本"] = (cg["投放金额"]/cg["客资量"].replace(0,pd.NA)).round(2)
        cg["成交成本"] = (cg["投放金额"]/cg["总成交量"].replace(0,pd.NA)).round(2)
        cg["成交率%"] = (cg["总成交量"]/cg["客资量"].replace(0,pd.NA)*100).round(2)
        st.dataframe(cg, use_container_width=True, hide_index=True)
        ca,cb,cc,cd = st.columns(4)
        with ca:
            fig = px.bar(cg,x="地区",y="客资量",title="客资量",color="地区",color_discrete_sequence=COLORS)
            st.plotly_chart(make_chart(fig),use_container_width=True)
        with cb:
            fig = px.bar(cg,x="地区",y="总成交量",title="成交量",color="地区",color_discrete_sequence=COLORS)
            st.plotly_chart(make_chart(fig),use_container_width=True)
        with cc:
            fig = px.bar(cg,x="地区",y="客资成本",title="客资成本",color="地区",color_discrete_sequence=COLORS)
            st.plotly_chart(make_chart(fig),use_container_width=True)
        with cd:
            fig = px.bar(cg,x="地区",y="成交成本",title="成交成本",color="地区",color_discrete_sequence=COLORS)
            st.plotly_chart(make_chart(fig),use_container_width=True)

with tab2:
    st.subheader("分渠道数据对比")
    if not df_detail_filtered.empty and "渠道分类" in df_detail_filtered.columns:
        rg = df_detail_filtered.groupby("渠道分类").agg(投放金额=("投放金额","sum"), 客资量=("客资量","sum"), 总成交量=("总成交量","sum")).reset_index()
        rg["客资成本"] = (rg["投放金额"]/rg["客资量"].replace(0,pd.NA)).round(2)
        rg = rg.sort_values("客资量", ascending=False)
        st.dataframe(rg, use_container_width=True, hide_index=True)
        ra,rb = st.columns(2)
        with ra:
            fig = px.bar(rg,x="渠道分类",y="客资量",title="各渠道客资量",color="渠道分类",color_discrete_sequence=COLORS)
            st.plotly_chart(make_chart(fig),use_container_width=True)
        with rb:
            fig = px.bar(rg,x="渠道分类",y="投放金额",title="各渠道投放金额",color="渠道分类",color_discrete_sequence=COLORS)
            st.plotly_chart(make_chart(fig),use_container_width=True)
        rc,rd = st.columns(2)
        with rc:
            fig = px.bar(rg,x="渠道分类",y="总成交量",title="各渠道成交量",color="渠道分类",color_discrete_sequence=COLORS)
            st.plotly_chart(make_chart(fig),use_container_width=True)
        with rd:
            fig = px.bar(rg,x="渠道分类",y="客资成本",title="各渠道客资成本",color="渠道分类",color_discrete_sequence=COLORS)
            st.plotly_chart(make_chart(fig),use_container_width=True)
        re,rf = st.columns(2)
        with re:
            fig = px.pie(rg,names="渠道分类",values="客资量",title="各渠道客资量占比",color_discrete_sequence=COLORS)
            fig.update_traces(textposition='inside', textinfo='percent+label')
            st.plotly_chart(make_chart(fig),use_container_width=True)
        with rf:
            fig = px.pie(rg,names="渠道分类",values="投放金额",title="各渠道投放占比",color_discrete_sequence=COLORS)
            fig.update_traces(textposition='inside', textinfo='percent+label')
            st.plotly_chart(make_chart(fig),use_container_width=True)
        st.divider()
        st.subheader("渠道月度趋势")
        df_ch_trend = apply_filter(df_detail, sel_city, "全部月份")
        if not df_ch_trend.empty and "渠道分类" in df_ch_trend.columns:
            ch_month = df_ch_trend.groupby(["月份","渠道分类"]).agg(客资量=("客资量","sum"), 投放金额=("投放金额","sum")).reset_index()
            ch_month["月份"] = pd.Categorical(ch_month["月份"], categories=MONTHS, ordered=True)
            ch_month = ch_month.sort_values("月份")
            fig = px.line(ch_month,x="月份",y="客资量",color="渠道分类",title="各渠道客资量月度趋势",markers=True,color_discrete_sequence=COLORS)
            st.plotly_chart(make_chart(fig),use_container_width=True)
            fig2 = px.line(ch_month,x="月份",y="投放金额",color="渠道分类",title="各渠道投放金额月度趋势",markers=True,color_discrete_sequence=COLORS)
            st.plotly_chart(make_chart(fig2),use_container_width=True)
    else:
        st.info("暂无渠道数据")

with tab3:
    st.subheader("趋势分析")
    df_trend = df_main.copy()
    if sel_city != "全部城市" and "地区" in df_trend.columns:
        df_trend = df_trend[df_trend["地区"] == sel_city]
    if not df_trend.empty and "月份" in df_trend.columns:
        cm_df = df_trend.groupby(["月份","地区"]).agg(投放金额=("投放金额","sum"), 客资量=("客资量","sum"), 总成交量=("总成交量","sum")).reset_index()
        cm_df["客资成本"] = (cm_df["投放金额"]/cm_df["客资量"].replace(0,pd.NA)).round(2)
        cm_df["成交成本"] = (cm_df["投放金额"]/cm_df["总成交量"].replace(0,pd.NA)).round(2)
        cm_df["月份"] = pd.Categorical(cm_df["月份"], categories=MONTHS, ordered=True)
        cm_df = cm_df.sort_values("月份")
        fig = px.line(cm_df,x="月份",y="客资成本",color="地区",title="各城市客资成本月度趋势",markers=True,color_discrete_sequence=COLORS)
        st.plotly_chart(make_chart(fig),use_container_width=True)
        fig2 = px.line(cm_df,x="月份",y="成交成本",color="地区",title="各城市成交成本月度趋势",markers=True,color_discrete_sequence=COLORS)
        st.plotly_chart(make_chart(fig2),use_container_width=True)
        tm = df_trend.groupby("月份").agg(客资量=("客资量","sum"), 总成交量=("总成交量","sum"), 投放金额=("投放金额","sum")).reset_index()
        tm["月份"] = pd.Categorical(tm["月份"], categories=MONTHS, ordered=True)
        tm = tm.sort_values("月份")
        fig3 = px.line(tm,x="月份",y=["客资量","总成交量"],title="月度客资与成交趋势",markers=True,color_discrete_sequence=COLORS)
        st.plotly_chart(make_chart(fig3),use_container_width=True)

with tab4:
    st.subheader("数据明细")
    t1,t2 = st.tabs(["四地汇总","渠道明细"])
    with t1:
        st.dataframe(df_filtered.dropna(axis=1,how='all'), use_container_width=True)
    with t2:
        st.dataframe(df_detail_filtered.dropna(axis=1,how='all'), use_container_width=True)
